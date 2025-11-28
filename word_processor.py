from docx import Document
import io

def extract_text_with_headers(file_bytes: bytes) -> str:
    try:
        doc = Document(io.BytesIO(file_bytes))
        full_text = []

        # 1. Extract from Headers (CRUCIAL for Resumes)
        for section in doc.sections:
            # Check standard header
            if section.header:
                for para in section.header.paragraphs:
                    if para.text.strip():
                        full_text.append(para.text.strip())
                # Headers can also contain tables
                for table in section.header.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                full_text.append(cell.text.strip())

        # 2. Extract from Main Body Paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())

        # 3. Extract from Main Body Tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    # distinct check to avoid duplicates if table text is included in paragraphs
                    # (Though usually tables are separate in python-docx structure)
                    for para in cell.paragraphs:
                        if para.text.strip():
                            full_text.append(para.text.strip())

        return "\n".join(full_text)

    except Exception as e:
        return f"Error: {e}"
